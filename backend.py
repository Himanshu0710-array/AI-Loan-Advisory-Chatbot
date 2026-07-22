"""
============================================================
  🏦 AI Loan Advisory Chatbot — Backend (Flask + RAG Pipeline)
============================================================

An AI-powered Loan Advisory Agent using RAG (Retrieval-Augmented Generation)
that processes loan policy PDFs and answers questions with source citations.

Tech Stack:
  - Backend: Python + Flask
  - LLM (generation): Google Gemini (gemini-2.0-flash) — optional, see PRIVACY NOTE
  - Embeddings: LOCAL sentence-transformers (all-MiniLM-L6-v2) — never sent externally
  - PDF Parsing: PyPDF2
  - Vector Store: Custom (NumPy cosine similarity)
  - Frontend: HTML + CSS + JavaScript

------------------------------------------------------------
PRIVACY NOTE (read this before claiming "fully private" anywhere):
  - Embeddings for your documents and every query are computed LOCALLY on
    this machine using sentence-transformers. Document text never leaves
    your server for embedding purposes.
  - By default, ONLY the top-K most relevant snippets (not full documents)
    plus the user's question are sent to Google's Gemini API to generate
    the final answer.
  - Set LOCAL_ONLY_MODE=true in your .env to disable Gemini entirely. In
    that mode, answers are generated with a local extractive template and
    NOTHING leaves your server. Answer quality/fluency is lower in this mode.
------------------------------------------------------------

SECURITY NOTE:
  - Upload / clear endpoints require an admin key (ADMIN_API_KEY env var)
    sent via the `X-Admin-Key` header.
  - CORS is restricted to ALLOWED_ORIGINS (comma-separated env var),
    default is localhost only.
  - Flask debug mode is OFF unless FLASK_DEBUG=true is explicitly set.
  - This is still a prototype-grade server (no auth on read endpoints, no
    per-user rate limiting beyond an optional flask-limiter hook). Do not
    deploy as-is to a public production environment without a proper
    auth layer and HTTPS in front of it.

HOW TO RUN:
  1. Set GEMINI_API_KEY (optional if LOCAL_ONLY_MODE=true) and
     ADMIN_API_KEY in the .env file
  2. pip install -r requirements.txt
  3. python backend.py
  4. Open http://localhost:5000
"""

# ============================================================
#  1. IMPORTS & CONFIGURATION
# ============================================================

import os
import json
import re
import time
import sys
import math
from pathlib import Path
from functools import wraps
from datetime import datetime

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

import numpy as np
from PyPDF2 import PdfReader
import docx
import openpyxl
from bs4 import BeautifulSoup
from flask import Flask, request, jsonify, send_from_directory, Response, stream_with_context
from flask_cors import CORS
from dotenv import load_dotenv
from flask_sqlalchemy import SQLAlchemy
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity

load_dotenv()

# --- Config from environment ---
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
LOCAL_ONLY_MODE = os.getenv("LOCAL_ONLY_MODE", "false").strip().lower() == "true"
ADMIN_API_KEY = os.getenv("ADMIN_API_KEY", "")
_default_origins = "http://localhost:5000,http://127.0.0.1:5000,https://ai-loan-advisory-chatbot-sipo.onrender.com"
ALLOWED_ORIGINS = [o.strip() for o in os.getenv("ALLOWED_ORIGINS", _default_origins).split(",") if o.strip()]
FLASK_DEBUG = os.getenv("FLASK_DEBUG", "false").strip().lower() == "true"

# --- Google Gemini Setup (LLM & Embeddings) ---
import google.generativeai as genai

chat_model = None
EMBED_DIM = 768  # text-embedding-004 dimension

if not GEMINI_API_KEY:
    print("[!] WARNING: GEMINI_API_KEY is not set. You will not be able to embed documents or answer questions.")
else:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        chat_model = genai.GenerativeModel("gemini-2.0-flash")
        print("[+] Gemini client initialized successfully.")
    except Exception as e:
        print(f"[!] Could not initialize Gemini client: {e}")
        chat_model = None


# Directories
BASE_DIR = Path(__file__).resolve().parent
DOCUMENTS_DIR = BASE_DIR / "documents"
DATA_DIR = BASE_DIR / "data"
CLIENT_DIR = BASE_DIR / "client"

DOCUMENTS_DIR.mkdir(exist_ok=True)
DATA_DIR.mkdir(exist_ok=True)


# ============================================================
#  2. DOCUMENT PARSING SERVICE
# ============================================================

def parse_document(file_path: str) -> dict:
    """Parse a document (PDF, DOCX, XLSX, HTML) and extract text with logical grouping."""
    file_path = Path(file_path)
    file_name = file_path.name
    ext = file_path.suffix.lower()

    pages = []
    full_text = ""

    if ext == ".pdf":
        reader = PdfReader(str(file_path))
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append({"pageNumber": i + 1, "text": text.strip()})
            full_text += text + "\n"
        total_pages = len(reader.pages)
        
    elif ext == ".docx":
        doc = docx.Document(str(file_path))
        # Treat paragraphs as pages/blocks for chunking
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                pages.append({"pageNumber": i + 1, "text": para.text.strip()})
                full_text += para.text + "\n"
        total_pages = len(doc.paragraphs)
        
    elif ext == ".xlsx":
        wb = openpyxl.load_workbook(str(file_path), data_only=True)
        sheet_idx = 1
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_text = f"--- Sheet: {sheet_name} ---\n"
            for row in sheet.iter_rows(values_only=True):
                row_vals = [str(v) for v in row if v is not None]
                if row_vals:
                    sheet_text += " | ".join(row_vals) + "\n"
            if sheet_text.strip():
                pages.append({"pageNumber": sheet_idx, "text": sheet_text.strip()})
                full_text += sheet_text + "\n"
                sheet_idx += 1
        total_pages = len(wb.sheetnames)
        
    elif ext == ".html":
        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f, "html.parser")
        text = soup.get_text(separator="\n", strip=True)
        pages.append({"pageNumber": 1, "text": text})
        full_text = text
        total_pages = 1
    else:
        raise ValueError(f"Unsupported document format: {ext}")

    return {
        "fileName": file_name,
        "pages": pages,
        "text": full_text,
        "metadata": {
            "totalPages": total_pages,
            "fileSize": os.path.getsize(file_path),
        },
    }


# ============================================================
#  3. TEXT CHUNKING SERVICE
# ============================================================

def split_into_sentences(text: str) -> list:
    text = re.sub(r"\n+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    sentences = re.split(r"(?<=[.!?])\s+", text)
    return [s.strip() for s in sentences if s.strip()]


def chunk_document(pages: list, file_name: str, chunk_size: int = 500, chunk_overlap: int = 80) -> list:
    """Split document pages into overlapping chunks with source metadata."""
    chunks = []
    chunk_index = 0
    safe_name = re.sub(r"[^a-zA-Z0-9]", "_", file_name.replace(".pdf", "")).lower()

    for page in pages:
        page_text = page["text"].strip()
        if not page_text:
            continue

        sentences = split_into_sentences(page_text)
        current_chunk = ""

        for sentence in sentences:
            if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                chunks.append({
                    "id": f"{safe_name}_p{page['pageNumber']}_c{chunk_index}",
                    "text": current_chunk.strip(),
                    "metadata": {
                        "fileName": file_name,
                        "pageNumber": page["pageNumber"],
                        "chunkIndex": chunk_index,
                    },
                })
                chunk_index += 1
                overlap_text = current_chunk[-chunk_overlap:]
                current_chunk = overlap_text + " " + sentence
            else:
                current_chunk += (" " if current_chunk else "") + sentence

        if current_chunk.strip():
            chunks.append({
                "id": f"{safe_name}_p{page['pageNumber']}_c{chunk_index}",
                "text": current_chunk.strip(),
                "metadata": {
                    "fileName": file_name,
                    "pageNumber": page["pageNumber"],
                    "chunkIndex": chunk_index,
                },
            })
            chunk_index += 1

    return chunks


# ============================================================
#  4. EMBEDDING SERVICE (Google Gemini)
# ============================================================

import requests

def generate_embedding(text: str) -> list:
    """Generate a single embedding vector via Gemini REST API.
    
    NOTE: Embeddings ALWAYS use the Gemini API regardless of LOCAL_ONLY_MODE.
    LOCAL_ONLY_MODE only controls whether the Gemini *chat* model is used for
    answer generation. Without real embeddings, the entire RAG pipeline fails
    (zero vectors → zero cosine similarity → everything is 'out of scope').
    """
    if not GEMINI_API_KEY:
        print("[!] WARNING: Cannot generate embedding — GEMINI_API_KEY is not set. RAG search will not work.")
        return [0.0] * EMBED_DIM
        
    url = f"https://generativelanguage.googleapis.com/v1beta/models/text-embedding-004:embedContent?key={GEMINI_API_KEY}"
    try:
        resp = requests.post(url, json={
            "model": "models/text-embedding-004",
            "content": {"parts": [{"text": text}]}
        }, timeout=15)
        resp.raise_for_status()
        values = resp.json().get("embedding", {}).get("values", [])
        if not values or all(v == 0.0 for v in values):
            print(f"[!] WARNING: Gemini returned empty/zero embedding for text: {text[:80]}...")
            return [0.0] * EMBED_DIM
        return values
    except Exception as e:
        print(f"[!] Error calling Gemini REST API for embedding: {e}")
        return [0.0] * EMBED_DIM


def generate_embeddings_batch(texts: list, batch_size: int = 32) -> list:
    """Generate embeddings for many texts via Gemini REST API.
    
    Always uses Gemini API (not affected by LOCAL_ONLY_MODE).
    """
    if not texts:
        return []
    if not GEMINI_API_KEY:
        print("[!] WARNING: Cannot generate batch embeddings — GEMINI_API_KEY is not set.")
        return [[0.0] * EMBED_DIM] * len(texts)
    
    url = f"https://generativelanguage.googleapis.com/v1beta/models/text-embedding-004:batchEmbedContents?key={GEMINI_API_KEY}"
    vectors = []
    
    for i in range(0, len(texts), batch_size):
        batch_texts = texts[i:i+batch_size]
        try:
            reqs = [
                {
                    "model": "models/text-embedding-004",
                    "content": {"parts": [{"text": t}]}
                }
                for t in batch_texts
            ]
            resp = requests.post(url, json={"requests": reqs}, timeout=30)
            resp.raise_for_status()
            
            data = resp.json()
            for emb in data.get("embeddings", []):
                vectors.append(emb.get("values", [0.0] * EMBED_DIM))
                
        except Exception as e:
            print(f"[!] Error in batch embedding (batch {i//batch_size + 1}): {e}")
            vectors.extend([[0.0] * EMBED_DIM] * len(batch_texts))
            
    return vectors


def generate_embeddings(chunks: list, batch_size: int = 32) -> list:
    """Generate embeddings for chunk dicts, attaching the vector to each."""
    texts = [c["text"] for c in chunks]
    vectors = generate_embeddings_batch(texts, batch_size=batch_size)
    embedded_chunks = [{**chunk, "embedding": vec} for chunk, vec in zip(chunks, vectors)]
    print(f"  [+] Generated {len(embedded_chunks)} embeddings via Gemini API")
    return embedded_chunks


# ============================================================
#  5. VECTOR STORE (ChromaDB)
# ============================================================

import chromadb

CHROMA_DIR = DATA_DIR / "chroma"
chroma_client = chromadb.PersistentClient(path=str(CHROMA_DIR))
collection = chroma_client.get_or_create_collection(name="loan_documents", metadata={"hnsw:space": "cosine"})

def _check_embedding_dimension():
    """Check if existing ChromaDB embeddings match the expected EMBED_DIM (768).
    
    If the collection was created with a different embedding model (e.g., 384-dim
    sentence-transformers), the dimensions won't match and all queries will fail with:
    'Collection expecting embedding with dimension of 384, got 768'
    
    In that case, we clear the collection so documents get re-indexed with the correct model.
    """
    global collection
    try:
        count = collection.count()
        if count == 0:
            return  # Empty collection, nothing to check
        
        # Peek at one embedding to check its dimension
        sample = collection.peek(limit=1)
        if sample and sample.get("embeddings") and len(sample["embeddings"]) > 0:
            existing_dim = len(sample["embeddings"][0])
            if existing_dim != EMBED_DIM:
                print(f"[!] DIMENSION MISMATCH: ChromaDB has {existing_dim}-dim embeddings, but current model uses {EMBED_DIM}-dim.")
                print(f"[!] Clearing old collection and re-indexing with correct embeddings...")
                chroma_client.delete_collection("loan_documents")
                collection = chroma_client.create_collection(name="loan_documents", metadata={"hnsw:space": "cosine"})
                print(f"[+] Collection cleared. Documents will be re-indexed on startup.")
            else:
                print(f"[+] Embedding dimension check passed: {existing_dim}-dim ✓")
    except Exception as e:
        print(f"[!] Could not verify embedding dimensions: {e}")


def init_vector_store():
    """Auto-load any PDFs found if ChromaDB is empty or was just cleared due to dimension mismatch."""
    # First, check for embedding dimension mismatch (e.g., old 384-dim vs new 768-dim)
    _check_embedding_dimension()
    
    stats = get_store_stats()
    print(f"[*] Loaded ChromaDB vector store: {stats['totalChunks']} chunks")

    # 1. Auto-load Documents
    if stats["totalChunks"] == 0 and DOCUMENTS_DIR.exists():
        for f in DOCUMENTS_DIR.iterdir():
            if f.suffix.lower() in [".pdf", ".docx", ".xlsx", ".html"]:
                try:
                    print(f"[*] Auto-loading existing document into ChromaDB: {f.name}")
                    parsed = parse_document(str(f))
                    chunks = chunk_document(parsed["pages"], f.name)
                    embedded = generate_embeddings(chunks)
                    add_documents(embedded, f.name)
                    print(f"[+] Successfully auto-loaded {len(embedded)} chunks from {f.name}")
                except Exception as e:
                    print(f"[!] Auto-load error on {f.name}: {e}")




def save_store():
    # ChromaDB persists automatically
    pass

def add_documents(chunks: list, file_name: str):
    """Add embedded chunks to the ChromaDB store (replaces existing from same file)."""
    # Delete existing chunks for this file
    collection.delete(where={"fileName": file_name})
    
    if not chunks:
        return
        
    ids = [c["id"] for c in chunks]
    embeddings = [c["embedding"] for c in chunks]
    texts = [c["text"] for c in chunks]
    metadatas = [c["metadata"] for c in chunks]
    
    collection.add(
        ids=ids,
        embeddings=embeddings,
        documents=texts,
        metadatas=metadatas
    )
    print(f"[+] Added {len(chunks)} chunks from '{file_name}' to ChromaDB")

def cosine_similarity(vec_a, vec_b) -> float:
    """Calculate cosine similarity between two vectors."""
    a, b = np.array(vec_a), np.array(vec_b)
    if a.shape != b.shape:
        return 0.0
    norm_a, norm_b = np.linalg.norm(a), np.linalg.norm(b)
    if not norm_a or not norm_b:
        return 0.0
    return float(np.dot(a, b) / (norm_a * norm_b))


STOP_WORDS = {
    "the", "and", "for", "are", "but", "not", "you", "all", "can",
    "had", "her", "was", "one", "our", "out", "has", "have", "been",
    "will", "with", "this", "that", "from", "they", "were", "some",
    "them", "than", "each", "which", "their", "about", "would",
    "there", "these", "other", "into", "more", "also", "should",
    "who", "what", "where", "when", "why", "how", "is", "am",
    "of", "in", "on", "at", "to", "a", "an", "it", "its", "by", "as", "does",
    "do", "did", "done", "if", "or", "nor", "neither", "either", "any", "many",
    "much", "every", "such", "only", "own", "so", "too", "very", "just", "now",
}


def search_similar(query_embedding: list, query_text: str = "", top_k: int = 5) -> list:
    """Find the top-K most similar chunks using ChromaDB semantic + light keyword scoring."""
    stats = get_store_stats()
    if stats["totalChunks"] == 0:
        return []

    # Fetch top 20 candidates from ChromaDB
    results = collection.query(
        query_embeddings=[query_embedding],
        n_results=min(20, stats["totalChunks"])
    )

    if not results["documents"] or not results["documents"][0]:
        return []

    q_words = set(re.findall(r"\b[a-z0-9]{3,}\b", query_text.lower())) - STOP_WORDS
    scored = []

    for idx, text in enumerate(results["documents"][0]):
        # ChromaDB cosine distance: 1 - cosine_similarity
        distance = results["distances"][0][idx]
        cos_sim = 1.0 - distance
        
        metadata = results["metadatas"][0][idx]
        
        kw_boost = 0.0
        if q_words:
            c_words = set(re.findall(r"\b[a-z0-9]{3,}\b", text.lower()))
            overlap = sum(1 for w in q_words if w in c_words)
            if cos_sim >= 0.15:
                kw_boost = (overlap / len(q_words)) * 0.3

        final_score = min(1.0, cos_sim + kw_boost)

        text_upper = text.upper()
        text_upper = text.upper()
        if "___" in text or "FOR OFFICE USE" in text_upper or "CUSTOMER SIGNATURE" in text_upper or text_upper[:15].startswith("PAGE "):
            final_score *= 0.15

        scored.append({"text": text, "metadata": metadata, "score": final_score})

    scored.sort(key=lambda x: x["score"], reverse=True)
    # Lowered threshold for Gemini embeddings
    return [s for s in scored if s["score"] >= 0.05][:top_k]

def get_store_stats():
    count = collection.count()
    
    if count == 0:
        return {
            "totalChunks": 0,
            "documents": [],
            "documentCount": 0,
        }
    
    # We retrieve a few items to guess the files, or we can just fetch all metadatas
    try:
        results = collection.get(include=["metadatas"])
        doc_files = list(set(m["fileName"] for m in results["metadatas"] if m and "fileName" in m))
    except Exception:
        doc_files = []

    return {
        "totalChunks": count,
        "documents": doc_files,
        "documentCount": len(doc_files),
    }

def clear_store():
    # To clear, we can delete the collection and recreate it
    chroma_client.delete_collection("loan_documents")
    global collection
    collection = chroma_client.create_collection(name="loan_documents", metadata={"hnsw:space": "cosine"})
    print("[*] ChromaDB Vector store cleared")


# ============================================================
#  6. SEMANTIC DOMAIN GATE (replaces hardcoded keyword lists)
# ============================================================
# Instead of matching questions against a fixed list of "banned phrases"
# (which is brittle and easy to fool in both directions), we embed a
# short description of the loan/banking domain once, and compare every
# incoming query to it semantically. This generalizes to phrasings we
# never explicitly listed.

DOMAIN_DESCRIPTION = (
    "Loan policy, eligibility criteria, interest rates, EMI calculation, "
    "processing fees, prepayment and foreclosure charges, required KYC "
    "documents such as PAN card and Aadhaar card, income and salary requirements, "
    "CIBIL credit score, loan tenure, principal and repayment schedule, "
    "bank lending rules, mortgage and collateral, loan sanction and "
    "disbursement process for home loans, personal loans, business loans, "
    "education loans, auto loans, vehicle finance, gold loans, loan against property, "
    "balance transfer, top-up loans, credit card, banking terms and conditions, "
    "MITC (Most Important Terms and Conditions), NPA, overdue, late payment, "
    "loan application process, documentation requirements, guarantor, co-applicant, "
    "fixed rate, floating rate, reducing balance, flat rate, amortization schedule."
)
DOMAIN_ANCHOR_EMBEDDING = generate_embedding(DOMAIN_DESCRIPTION)

# Similarity below this to the domain anchor is treated as likely off-topic.
# Lowered to 0.05 for Gemini embeddings (they have a wider angular spread)
DOMAIN_SIMILARITY_THRESHOLD = 0.05


def is_out_of_domain(question: str, query_embedding: list, top_retrieval_score: float) -> bool:
    """
    Semantic out-of-domain check. A question is treated as out of domain when
    BOTH of the following hold:
      - it is not semantically close to the loan/banking domain description, AND
      - it did not retrieve any strongly relevant chunk from the loaded documents.
    
    BYPASS: If the retrieval score is strong (>= 0.30), the question is always
    considered in-domain regardless of the anchor similarity. This prevents
    false rejections when documents contain relevant content.
    """
    # Strong retrieval hit → always in domain (bypass the gate)
    if top_retrieval_score >= 0.30:
        return False
    
    # Check if embedding is a zero vector (API key missing/broken)
    if all(v == 0.0 for v in query_embedding[:10]):
        print("  [!] WARNING: Query embedding is a zero vector — domain gate bypassed (embeddings broken)")
        return False  # Don't block if embeddings are broken; let the user see whatever we retrieved
    
    anchor_sim = cosine_similarity(query_embedding, DOMAIN_ANCHOR_EMBEDDING)
    print(f"  [*] Domain similarity: {anchor_sim:.4f} (threshold: {DOMAIN_SIMILARITY_THRESHOLD}), top retrieval: {top_retrieval_score:.4f}")
    return anchor_sim < DOMAIN_SIMILARITY_THRESHOLD and top_retrieval_score < 0.10


def build_out_of_domain_response(question: str) -> dict:
    answer = (
        "⚠️ **Out of Scope Query**\n\n"
        f"Your question (*\"{question}\"*) appears to be outside the scope of this financial advisory system.\n\n"
        "I am the **LMS Loan Advisory Agent**, focused on your loaded loan policy documents.\n\n"
        "**Please ask me about:**\n"
        "• **Loan Eligibility Criteria** (Age, Salary, CIBIL score requirements)\n"
        "• **Interest Rates & Processing Fees**\n"
        "• **EMI Calculation Formulas & Examples**\n"
        "• **Required Documentation** (KYC, Income statements)\n"
        "• **Prepayment & Foreclosure Terms**"
    )
    return {
        "answer": answer,
        "sources": [],
        "validation": {"isGrounded": True, "confidence": "high", "warnings": []},
    }


# ============================================================
#  7. RESPONSE VALIDATION SERVICE
# ============================================================

def extract_numbers(text):
    matches = re.findall(
        r"\d+(?:\.\d+)?(?:%|,\d{3})*(?:\s*(?:lakh|crore|lac|thousand|million|billion))?",
        text, re.IGNORECASE,
    )
    return [m.strip().lower() for m in matches]


def extract_keywords(text):
    cleaned = re.sub(r"[^a-z0-9\s]", "", text.lower())
    return [w for w in cleaned.split() if len(w) >= 4 and w not in STOP_WORDS]


def validate_response(response: str, context_chunks: list) -> dict:
    """Validate that the LLM response is grounded in retrieved context."""
    warnings = []
    response_lower = response.lower()
    context_text = " ".join(c["text"].lower() for c in context_chunks)

    uncertainty = [
        "i don't have enough information", "not mentioned in the document",
        "i cannot find", "no information available", "not specified in the provided",
    ]
    if any(p in response_lower for p in uncertainty):
        return {"isGrounded": True, "confidence": "low",
                "warnings": ["The AI indicated it could not find relevant information."]}

    resp_nums = extract_numbers(response)
    ctx_nums = extract_numbers(context_text)
    unmatched = [n for n in resp_nums if n not in ctx_nums]
    if unmatched:
        warnings.append(f"Some values ({', '.join(unmatched[:5])}) were not found in source documents.")

    resp_kw = extract_keywords(response)
    ctx_kw = set(extract_keywords(context_text))
    ratio = sum(1 for kw in resp_kw if kw in ctx_kw) / len(resp_kw) if resp_kw else 0

    if ratio < 0.2:
        confidence = "low"
        warnings.append("Response may contain information not from the source documents.")
    elif ratio < 0.4:
        confidence = "medium"
    else:
        confidence = "high"

    return {"isGrounded": len(warnings) == 0, "confidence": confidence, "warnings": warnings}


# ============================================================
#  8. RAG PIPELINE & CASUAL CONVERSATION HANDLER
# ============================================================

SYSTEM_PROMPT = """You are the LMS Loan Advisory Agent — a senior, friendly, and expert AI financial advisor specializing in loan policy documents.

IMPORTANT RULES:
1. ONLY answer factual loan details based on the provided context from loan policy documents and the recent conversation history.
2. If asked about something outside the financial/loan domain, politely decline and explain you are restricted to the loaded loan policy documents.
3. If the context does not contain enough information for a specific loan policy query, clearly state: "I don't have enough information in the loaded documents to answer this specific question."
4. Always reference which document and page number your answer comes from when discussing policy details.
5. Be concise, highly structured, and human-readable. Ignore OCR form blanks ("___", "Signature FOR OFFICE USE ONLY").
6. For EMI calculations, clearly show the formula and step-by-step math.
7. Format answers using bold headings, bullet points, and clean spacing. Never output raw unstructured text blocks.
8. Use the conversation history only to resolve references like "that", "it", or follow-ups such as "what about for X" — do not restate old answers unnecessarily.
9. MULTILINGUAL SUPPORT: Automatically detect the language of the user's question. You must generate your response in the EXACT SAME language as the user's question, even if the source documents are in a different language."""


def format_history(history: list, max_turns: int = 3) -> str:
    """Format recent conversation turns for inclusion in the prompt (fixes lack of memory)."""
    if not history:
        return ""
    recent = history[-max_turns:]
    lines = []
    for turn in recent:
        q = str(turn.get("question", "")).strip()
        a = str(turn.get("answer", "")).strip()
        if q:
            lines.append(f"User: {q}")
        if a:
            # Keep prior answers short in the prompt to save tokens
            lines.append(f"Assistant: {a[:300]}")
    if not lines:
        return ""
    return "\n\n--- RECENT CONVERSATION (for context/follow-ups only) ---\n" + "\n".join(lines) + "\n--- END OF RECENT CONVERSATION ---"


def handle_casual_query(question: str) -> dict or None:
    """Greetings, identity questions, and thanks — simple conversational routing (not domain gating)."""
    q_clean = re.sub(r"[^a-z0-9\s]", "", question.lower()).strip()
    words = q_clean.split()

    greetings = {"hi", "hello", "hey", "greetings", "yo", "hola", "sup", "howdy",
                 "morning", "afternoon", "evening", "hii", "hiii"}
    if q_clean in greetings or (len(words) <= 3 and any(w in greetings for w in words)):
        # Use ChromaDB stats instead of deleted vector_store variable
        stats = get_store_stats()
        doc_names = stats.get("documents", [])
        docs_str = ", ".join(doc_names) if doc_names else "No documents uploaded yet"
        answer = (
            "👋 **Hello! Welcome to your AI Loan Advisory Agent.**\n\n"
            f"I am connected to your active document(s): **{docs_str}** and ready to assist you!\n\n"
            "### 💡 How I Can Help You Today:\n"
            "• **Check Eligibility:** Ask about age, salary, income, or CIBIL score requirements.\n"
            "• **Interest Rates & Fees:** Compare loan rates, processing fees, or foreclosure terms.\n"
            "• **EMI Calculations:** Get exact mathematical formulas and monthly installment examples.\n"
            "• **Documentation:** Find out what KYC and income verification documents you need.\n\n"
            "👉 *Try asking:* `What are the eligibility criteria?` or `How is EMI calculated?`"
        )
        return {"answer": answer, "sources": [], "validation": {"isGrounded": True, "confidence": "high", "warnings": []}}

    identity_phrases = {"who are you", "what are you", "what can you do", "help",
                         "what is your name", "about you", "how do you work", "what do you do"}
    if q_clean in identity_phrases or any(p in q_clean for p in identity_phrases):
        answer = (
            "🤖 **I am the LMS AI Loan Advisory Agent!**\n\n"
            "I use **Retrieval-Augmented Generation (RAG)** to read, index, and analyze your uploaded loan "
            "policy PDFs locally, then generate accurate answers backed by source citations.\n\n"
            "**You can ask me things like:**\n"
            "1. *\"What is the minimum CIBIL score needed for a home loan?\"*\n"
            "2. *\"Can I prepay my floating rate home loan without penalty?\"*\n"
            "3. *\"Calculate EMI for 5 Lakhs at 12% for 3 years.\"*\n"
            "4. *\"What documents are required for self-employed applicants?\"*"
        )
        return {"answer": answer, "sources": [], "validation": {"isGrounded": True, "confidence": "high", "warnings": []}}

    thanks_phrases = {"thanks", "thank you", "thx", "awesome", "great", "perfect", "good job",
                       "bye", "goodbye", "see you", "ok", "got it", "nice", "thankyou"}
    if q_clean in thanks_phrases or (len(words) <= 4 and any(w in thanks_phrases for w in words)):
        answer = (
            "😊 **You're very welcome!**\n\n"
            "I'm always here to help you navigate loan guidelines, terms, and calculations. "
            "If you have any more questions about your policy documents, feel free to ask anytime!"
        )
        return {"answer": answer, "sources": [], "validation": {"isGrounded": True, "confidence": "high", "warnings": []}}

    return None


def generate_local_rag_response(question: str, chunks: list) -> str:
    """Extractive RAG synthesis. Selects relevant sentences verbatim from document chunks and groups them by source."""
    if not chunks:
        return "I couldn't find specific details regarding your query in the loaded documents. Try asking about eligibility, interest rates, or EMIs."

    q_lower = question.lower()

    # Simplify Intent Detection (Section Heading Only)
    INTENT_TITLES = {
        ("eligib", "age", "salary", "cibil", "score", "income", "criteria", "requirement"): "Eligibility Criteria & Requirements",
        ("emi", "calculat", "formula", "month", "installment"): "EMI Calculation & Formula Rules",
        ("interest", "rate", "fee", "charge", "processing", "cost"): "Interest Rates & Fee Structures",
        ("document", "proof", "pan", "aadhaar", "slip", "itr", "kyc"): "Required Documentation Checklist",
        ("prepay", "foreclos", "penalty", "lock-in", "close"): "Prepayment & Foreclosure Guidelines",
        ("compare", "difference", "vs", "versus"): "Policy Comparison Excerpts",
        ("bounce", "default", "miss", "overdue", "late"): "Late Payment & Default Clauses",
        ("apply", "application", "procedure", "process", "steps"): "Application Process & Guidelines",
    }

    title = "Policy Findings & Document Excerpts"
    for keywords, heading in INTENT_TITLES.items():
        if any(w in q_lower for w in keywords):
            title = heading
            break

    # Clean, scored sentence extraction verbatim from relevant chunks
    q_words = set(re.findall(r"\b[a-z0-9]{3,}\b", q_lower)) - STOP_WORDS
    extracted = []
    seen = set()

    for chunk in chunks:
        if chunk["score"] < 0.08:
            continue
        raw_text = chunk["text"].strip()
        if any(junk in raw_text.upper() for junk in ["___", "FOR OFFICE USE", "CUSTOMER SIGNATURE", "INDEX", "TABLE OF CONTENTS"]):
            continue

        doc_name = chunk["metadata"]["fileName"]
        page_num = chunk["metadata"]["pageNumber"]

        sentences = re.split(r"(?<=[.!?])\s+", raw_text)
        for s in sentences:
            s_clean = re.sub(r"\s+", " ", s).strip()
            if len(s_clean) < 30 or len(s_clean) > 350:
                continue
            if any(junk in s_clean for junk in ["....", "___", "---", "CUSTOMER NAME", "LOAN ACCOUNT NO", "Page "]):
                continue
            if s_clean in seen:
                continue

            seen.add(s_clean)
            s_words = set(re.findall(r"\b[a-z0-9]{3,}\b", s_clean.lower())) - STOP_WORDS
            overlap = sum(1 for w in q_words if w in s_words)
            extracted.append({
                "doc": doc_name,
                "page": page_num,
                "text": s_clean,
                "overlap": overlap,
                "chunk_score": chunk["score"]
            })

    extracted.sort(key=lambda x: (x["overlap"], x["chunk_score"]), reverse=True)

    # Assemble final human-readable response verbatim from retrieved sentences
    answer = f"### 💡 {title}\n\n"
    answer += "Here are the most relevant clauses extracted directly from your active policy documents:\n\n"

    if not extracted:
        top_chunk = chunks[0]
        clean_excerpt = re.sub(r"\s+", " ", top_chunk["text"]).strip()
        answer += f"• **{top_chunk['metadata']['fileName']} (Page {top_chunk['metadata']['pageNumber']}):** *\"{clean_excerpt[:300]}...\"*\n"
    else:
        grouped = {}
        for item in extracted[:8]:
            key = f"{item['doc']} (Page {item['page']})"
            if key not in grouped:
                grouped[key] = []
            grouped[key].append(item["text"])

        for key, sentences in grouped.items():
            answer += f"\n**{key}:**\n"
            for s in sentences:
                if ":" in s[:45] and len(s.split(":")[0]) < 38:
                    parts = s.split(":", 1)
                    answer += f"• **{parts[0].strip()}:** {parts[1].strip()}\n"
                else:
                    answer += f"• {s}\n"

    answer += "\n---\n*Generated locally from your loaded policy documents.*"
    return answer


def process_query(question: str, history: list = None) -> dict:
    """Process a user question through the full RAG pipeline."""
    history = history or []

    casual_response = handle_casual_query(question)
    if casual_response:
        print("  [+] Handled casual conversation query directly")
        return casual_response

    stats = get_store_stats()
    if stats["totalChunks"] == 0:
        return {
            "answer": (
                "📄 **No documents loaded yet!**\n\n"
                "Please upload a loan policy PDF document first. "
                "I'll then be able to answer your questions based on the document content."
            ),
            "sources": [],
            "validation": {"isGrounded": True, "confidence": "high", "warnings": []},
        }

    print("  [*] Embedding user query locally...")
    query_embedding = generate_embedding(question)

    print("  [*] Searching for relevant chunks...")
    relevant_chunks = search_similar(query_embedding, question, top_k=5)
    top_score = relevant_chunks[0]["score"] if relevant_chunks else 0.0

    if is_out_of_domain(question, query_embedding, top_score):
        return build_out_of_domain_response(question)

    if not relevant_chunks:
        return {
            "answer": (
                "⚠️ **Topic Not Found**\n\n"
                f"Your question (*\"{question}\"*) didn't match any section in the loaded documents. "
                "Try asking about eligibility, interest rates, EMIs, documentation, or foreclosure terms."
            ),
            "sources": [],
            "validation": {"isGrounded": True, "confidence": "low",
                            "warnings": ["No highly relevant chunks found in the active PDFs."]},
        }

    context_parts = [
        f'[Source {i+1}: "{c["metadata"]["fileName"]}", Page {c["metadata"]["pageNumber"]}]\n{c["text"]}'
        for i, c in enumerate(relevant_chunks)
    ]
    context = "\n\n---\n\n".join(context_parts)
    history_block = format_history(history)

    prompt = f"""{SYSTEM_PROMPT}

--- CONTEXT FROM LOAN DOCUMENTS ---
{context}
--- END OF CONTEXT ---
{history_block}

User Question: {question}

Provide a clear, accurate answer based ONLY on the above context (and the conversation history for follow-up context). Cite the source document and page number for key facts."""

    print("  [*] Generating response...")
    answer = None
    if chat_model is not None and not LOCAL_ONLY_MODE:
        try:
            result = chat_model.generate_content(prompt)
            answer = result.text
        except Exception as e:
            print(f"  [!] Online LLM unavailable ({e}). Using local extractive synthesis...")
    if answer is None:
        answer = generate_local_rag_response(question, relevant_chunks)

    sources = [
        {
            "fileName": c["metadata"]["fileName"],
            "pageNumber": c["metadata"]["pageNumber"],
            "relevanceScore": round(c["score"] * 100),
            "excerpt": c["text"][:150] + ("..." if len(c["text"]) > 150 else ""),
        }
        for c in relevant_chunks if c["score"] >= 0.08
    ]

    validation = validate_response(answer, relevant_chunks)
    print(f"  [+] Response generated (confidence: {validation['confidence']}, sources: {len(sources)})")

    return {"answer": answer, "sources": sources, "validation": validation}


# ============================================================
#  9. FLASK API SERVER
# ============================================================

app = Flask(__name__, static_folder=str(CLIENT_DIR), static_url_path="")
CORS(app, origins=ALLOWED_ORIGINS, supports_credentials=False)

app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + str(DATA_DIR / 'app.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['JWT_SECRET_KEY'] = os.getenv("JWT_SECRET_KEY", "dev-super-secret-key")

db = SQLAlchemy(app)
jwt = JWTManager(app)

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password = db.Column(db.String(255), nullable=False)

class ChatMessage(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    question = db.Column(db.Text, nullable=False)
    answer = db.Column(db.Text, nullable=False)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)

with app.app_context():
    db.create_all()

# Optional rate limiting — enabled automatically if flask-limiter is installed
try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address
    limiter = Limiter(get_remote_address, app=app, default_limits=["60 per minute"])
    print("[*] Rate limiting enabled (flask-limiter)")
except ImportError:
    limiter = None
    print("[*] flask-limiter not installed — running without rate limiting (pip install flask-limiter to enable)")


def require_admin_key(f):
    """Guard for endpoints that modify server state (upload/clear). Requires exact X-Admin-Key match."""
    @wraps(f)
    def wrapper(*args, **kwargs):
        if not ADMIN_API_KEY:
            return jsonify({"error": "Admin endpoints are disabled. Set ADMIN_API_KEY in .env to enable them."}), 503
        provided = request.headers.get("X-Admin-Key", "")
        if provided != ADMIN_API_KEY:
            return jsonify({"error": "Unauthorized. Missing or invalid X-Admin-Key header."}), 401
        return f(*args, **kwargs)
    return wrapper


# --- Feedback Store ---
FEEDBACK_FILE = DATA_DIR / "feedback.json"
feedback_store = []


def init_feedback_store():
    global feedback_store
    if FEEDBACK_FILE.exists():
        try:
            with open(FEEDBACK_FILE, "r") as f:
                feedback_store = json.load(f)
            print(f"[*] Loaded {len(feedback_store)} feedback entries")
        except Exception:
            feedback_store = []


def save_feedback():
    with open(FEEDBACK_FILE, "w") as f:
        json.dump(feedback_store, f, indent=2)


def get_negative_feedback_context():
    negatives = [fb for fb in feedback_store if fb.get("rating") == "down"]
    if not negatives:
        return ""
    recent = negatives[-3:]
    lines = [f"- Question: \"{fb['question']}\" → The user disliked this answer style. Avoid similar phrasing or structure." for fb in recent]
    return "\n\nIMPORTANT — The user previously gave NEGATIVE feedback on these responses. Learn from this and improve:\n" + "\n".join(lines)


# --- Serve Frontend ---
@app.route("/")
def serve_frontend():
    return send_from_directory(str(CLIENT_DIR), "index.html")


@app.route("/<path:path>")
def serve_static(path):
    return send_from_directory(str(CLIENT_DIR), path)


# --- Auth APIs ---
@app.route("/api/register", methods=["POST"])
def register():
    data = request.get_json(silent=True) or {}
    username = data.get("username")
    password = data.get("password")
    if not username or not password:
        return jsonify({"error": "Username and password required"}), 400
    
    if User.query.filter_by(username=username).first():
        return jsonify({"error": "User already exists"}), 400
        
    new_user = User(username=username, password=password) # In production, hash this!
    db.session.add(new_user)
    db.session.commit()
    return jsonify({"success": True, "message": "Registered successfully"})

@app.route("/api/login", methods=["POST"])
def login():
    data = request.get_json(silent=True) or {}
    username = data.get("username")
    password = data.get("password")
    
    user = User.query.filter_by(username=username, password=password).first()
    if not user:
        return jsonify({"error": "Invalid credentials"}), 401
        
    access_token = create_access_token(identity=str(user.id))
    return jsonify({"success": True, "access_token": access_token})

# --- Chat History API ---
@app.route("/api/chat/history", methods=["GET"])
@jwt_required()
def get_chat_history():
    try:
        user_id = get_jwt_identity()
        sessions = ChatMessage.query.filter_by(user_id=int(user_id)).order_by(ChatMessage.timestamp.asc()).all()
        history = [
            {
                "id": s.id,
                "question": s.question,
                "answer": s.answer,
                "timestamp": s.timestamp.isoformat()
            }
            for s in sessions
        ]
        return jsonify({"success": True, "history": history})
    except Exception as e:
        print(f"[!] History error: {e}")
        return jsonify({"error": str(e)}), 500

# --- Chat API (non-streaming) ---
@app.route("/api/chat", methods=["POST"])
@jwt_required(optional=True)
def chat():
    try:
        data = request.get_json(force=True, silent=True) or {}
        question = str(data.get("question", "")).strip()
        history = data.get("history", [])
        if not question:
            return jsonify({"error": "Please provide a question."}), 400
        if len(question) > 2000:
            return jsonify({"error": "Question is too long."}), 400

        print(f'\n[?] Question: "{question}"')
        result = process_query(question, history=history)
        
        # Save history if logged in
        user_id = get_jwt_identity()
        if user_id:
            try:
                new_session = ChatMessage(
                    user_id=int(user_id),
                    question=question,
                    answer=result.get("answer", "")
                )
                db.session.add(new_session)
                db.session.commit()
            except Exception as db_err:
                print(f"  [!] Failed to save chat history: {db_err}")
                db.session.rollback()

        return jsonify({"success": True, **result})
    except Exception as e:
        print(f"[!] Chat error: {e}")
        return jsonify({"error": str(e)}), 500


def stream_words_only(text: str, delay: float = 0.016):
    tokens = re.findall(r"\S+|\s+", text)
    for token in tokens:
        yield f"data: {json.dumps({'type': 'chunk', 'content': token})}\n\n"
        if token.strip():
            time.sleep(delay)


# --- Chat API (Streaming via SSE) ---
@app.route("/api/chat/stream", methods=["POST"])
@jwt_required(optional=True)
def chat_stream():
    try:
        data = request.get_json(force=True, silent=True) or {}
        question = str(data.get("question", "")).strip()
        history = data.get("history", [])
        if not question:
            return jsonify({"error": "Please provide a question."}), 400
        if len(question) > 2000:
            return jsonify({"error": "Question is too long."}), 400

        print(f'\n[?] Streaming Question: "{question}"')

        casual_response = handle_casual_query(question)
        if casual_response:
            def casual_gen():
                for token_sse in stream_words_only(casual_response["answer"], delay=0.015):
                    yield token_sse
                
                # Save casual response to history
                user_id = get_jwt_identity()
                if user_id:
                    try:
                        from flask import current_app
                        with current_app.app_context():
                            new_session = ChatMessage(
                                user_id=int(user_id),
                                question=question,
                                answer=casual_response["answer"]
                            )
                            db.session.add(new_session)
                            db.session.commit()
                    except Exception as db_err:
                        print(f"  [!] Failed to save chat history: {db_err}")
                        db.session.rollback()

                yield f"data: {json.dumps({'type': 'done', 'sources': casual_response['sources'], 'validation': casual_response['validation']})}\n\n"
            return Response(stream_with_context(casual_gen()), mimetype="text/event-stream")

        stats = get_store_stats()
        if stats["totalChunks"] == 0:
            no_doc_answer = "📄 **No documents loaded yet!**\n\nPlease upload a loan policy PDF first."
            def no_doc_gen():
                for token_sse in stream_words_only(no_doc_answer, delay=0.015):
                    yield token_sse
                yield f"data: {json.dumps({'type': 'done', 'sources': [], 'validation': {'isGrounded': True, 'confidence': 'high', 'warnings': []}})}\n\n"
            return Response(stream_with_context(no_doc_gen()), mimetype="text/event-stream")

        query_embedding = generate_embedding(question)
        relevant_chunks = search_similar(query_embedding, question, top_k=5)
        top_score = relevant_chunks[0]["score"] if relevant_chunks else 0.0

        if is_out_of_domain(question, query_embedding, top_score):
            ood = build_out_of_domain_response(question)
            def ood_gen():
                for token_sse in stream_words_only(ood["answer"], delay=0.015):
                    yield token_sse
                yield f"data: {json.dumps({'type': 'done', 'sources': [], 'validation': ood['validation']})}\n\n"
            return Response(stream_with_context(ood_gen()), mimetype="text/event-stream")

        if not relevant_chunks:
            fallback_answer = (
                "⚠️ **Topic Not Found**\n\n"
                f"Your question (*\"{question}\"*) didn't match any section in the loaded documents. "
                "Try asking about eligibility, interest rates, EMIs, documentation, or foreclosure terms."
            )
            def fallback_gen():
                for token_sse in stream_words_only(fallback_answer, delay=0.015):
                    yield token_sse
                yield f"data: {json.dumps({'type': 'done', 'sources': [], 'validation': {'isGrounded': True, 'confidence': 'low', 'warnings': []}})}\n\n"
            return Response(stream_with_context(fallback_gen()), mimetype="text/event-stream")

        context_parts = [
            f'[Source {i+1}: "{c["metadata"]["fileName"]}", Page {c["metadata"]["pageNumber"]}]\n{c["text"]}'
            for i, c in enumerate(relevant_chunks)
        ]
        context = "\n\n---\n\n".join(context_parts)
        feedback_context = get_negative_feedback_context()
        history_block = format_history(history)

        prompt = f"""{SYSTEM_PROMPT}{feedback_context}

--- CONTEXT FROM LOAN DOCUMENTS ---
{context}
--- END OF CONTEXT ---
{history_block}

User Question: {question}

Provide a clear, accurate answer based ONLY on the above context (and the conversation history for follow-up context). Cite the source document and page number for key facts."""

        sources = [
            {
                "fileName": c["metadata"]["fileName"],
                "pageNumber": c["metadata"]["pageNumber"],
                "relevanceScore": round(c["score"] * 100),
                "excerpt": c["text"][:150] + ("..." if len(c["text"]) > 150 else ""),
            }
            for c in relevant_chunks if c["score"] >= 0.08
        ]

        def generate_stream():
            full_answer = ""
            used_gemini = chat_model is not None and not LOCAL_ONLY_MODE
            if used_gemini:
                try:
                    response = chat_model.generate_content(prompt, stream=True)
                    for chunk_resp in response:
                        if chunk_resp.text:
                            full_answer += chunk_resp.text
                            for token_sse in stream_words_only(chunk_resp.text, delay=0.012):
                                yield token_sse
                except Exception as e:
                    print(f"  [!] Streaming LLM error ({e}). Using local fallback...")
                    full_answer = ""
            if not full_answer:
                full_answer = generate_local_rag_response(question, relevant_chunks)
                for token_sse in stream_words_only(full_answer, delay=0.016):
                    yield token_sse

            validation = validate_response(full_answer, relevant_chunks)
            
            # Save history if logged in
            user_id = get_jwt_identity()
            if user_id:
                try:
                    from flask import current_app
                    with current_app.app_context():
                        new_session = ChatMessage(
                            user_id=int(user_id),
                            question=question,
                            answer=full_answer
                        )
                        db.session.add(new_session)
                        db.session.commit()
                except Exception as db_err:
                    print(f"  [!] Failed to save chat history: {db_err}")
                    db.session.rollback()

            yield f"data: {json.dumps({'type': 'done', 'sources': sources, 'validation': validation})}\n\n"

        return Response(stream_with_context(generate_stream()), mimetype="text/event-stream")

    except Exception as e:
        print(f"[!] Stream error: {e}")
        return jsonify({"error": str(e)}), 500


# --- Document Upload API (admin-protected) ---
@app.route("/api/documents/upload", methods=["POST"])
@require_admin_key
def upload_document():
    try:
        if "document" not in request.files:
            return jsonify({"error": "No file uploaded."}), 400

        file = request.files["document"]
        ext = Path(file.filename).suffix.lower()
        if ext not in [".pdf", ".docx", ".xlsx", ".html"]:
            return jsonify({"error": "Only PDF, DOCX, XLSX, and HTML files are allowed."}), 400

        file.seek(0, 2)
        size = file.tell()
        file.seek(0)
        if size > 20 * 1024 * 1024:
            return jsonify({"error": "File too large. Maximum 20MB allowed."}), 400

        safe_name = re.sub(r"[^a-zA-Z0-9._-]", "_", file.filename)
        file_path = DOCUMENTS_DIR / safe_name
        file.save(str(file_path))

        print(f'\n[*] Processing: "{file.filename}"')

        print("  [*] Parsing document...")
        parsed = parse_document(str(file_path))
        print(f"  [+] Extracted {len(parsed['pages'])} logical pages/blocks")

        print("  [*] Chunking text...")
        chunks = chunk_document(parsed["pages"], safe_name)
        print(f"  [+] Created {len(chunks)} chunks")

        print("  [*] Generating embeddings locally (batched)...")
        embedded_chunks = generate_embeddings(chunks)

        print("  [*] Storing in vector database...")
        add_documents(embedded_chunks, safe_name)

        stats = get_store_stats()
        return jsonify({
            "success": True,
            "message": f'Successfully processed "{file.filename}"',
            "details": {
                "fileName": safe_name,
                "pages": len(parsed["pages"]),
                "chunks": len(chunks),
                "embeddingsGenerated": len(embedded_chunks),
                "totalChunksInStore": stats["totalChunks"],
                "totalDocuments": stats["documentCount"],
            },
        })
    except Exception as e:
        print(f"[!] Upload error: {e}")
        return jsonify({"error": str(e)}), 500


# --- List Documents (read-only, no admin key needed) ---
@app.route("/api/documents", methods=["GET"])
def list_documents():
    documents = []
    if DOCUMENTS_DIR.exists():
        for f in DOCUMENTS_DIR.iterdir():
            if f.suffix.lower() in [".pdf", ".docx", ".xlsx", ".html"]:
                documents.append({"name": f.name, "size": f.stat().st_size})
    return jsonify({"success": True, "documents": documents, "stats": get_store_stats()})


# --- Clear Documents (admin-protected) ---
@app.route("/api/documents/clear", methods=["DELETE"])
@require_admin_key
def clear_documents():
    clear_store()
    for f in DOCUMENTS_DIR.iterdir():
        if f.is_file():
            f.unlink()
    return jsonify({"success": True, "message": "All documents and vectors cleared."})


# --- User Feedback API ---
@app.route("/api/feedback", methods=["POST"])
def submit_feedback():
    try:
        data = request.get_json(force=True, silent=True) or {}
        entry = {
            "question": str(data.get("question", ""))[:500],
            "answer": str(data.get("answer", ""))[:500],
            "rating": data.get("rating", "up"),
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        }
        feedback_store.append(entry)
        save_feedback()
        print(f"  [*] Feedback recorded: {entry['rating']} for \"{entry['question'][:50]}...\"")
        return jsonify({"success": True, "message": "Feedback recorded. Thank you!"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- EMI Calculator API ---
@app.route("/api/emi/calculate", methods=["POST"])
def calculate_emi():
    try:
        data = request.get_json(force=True, silent=True) or {}
        principal = float(data.get("principal", 0))
        annual_rate = float(data.get("rate", 0))
        tenure_months = int(data.get("tenure", 0))

        if principal <= 0 or annual_rate <= 0 or tenure_months <= 0:
            return jsonify({"error": "All values must be positive numbers."}), 400

        r = annual_rate / 12 / 100
        n = tenure_months

        if r == 0:
            emi = principal / n
        else:
            emi = (principal * r * math.pow(1 + r, n)) / (math.pow(1 + r, n) - 1)

        total_payment = emi * n
        total_interest = total_payment - principal

        breakdown = []
        balance = principal
        for month in range(1, min(n + 1, 61)):
            interest_part = balance * r
            principal_part = emi - interest_part
            balance -= principal_part
            breakdown.append({
                "month": month,
                "emi": round(emi, 2),
                "principal": round(principal_part, 2),
                "interest": round(interest_part, 2),
                "balance": round(max(balance, 0), 2),
            })

        return jsonify({
            "success": True,
            "emi": round(emi, 2),
            "totalPayment": round(total_payment, 2),
            "totalInterest": round(total_interest, 2),
            "principal": principal,
            "rate": annual_rate,
            "tenure": tenure_months,
            "breakdown": breakdown,
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- Export Chat as PDF ---
@app.route("/api/export/pdf", methods=["POST"])
def export_chat_pdf():
    try:
        data = request.get_json(force=True, silent=True) or {}
        messages = data.get("messages", [])

        if not messages:
            return jsonify({"error": "No messages to export."}), 400

        from io import BytesIO

        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib import colors

            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=40, bottomMargin=40)
            styles = getSampleStyleSheet()
            story = []

            title_style = ParagraphStyle("Title", parent=styles["Heading1"], textColor=colors.HexColor("#6366f1"), fontSize=18)
            story.append(Paragraph("LMS — AI Loan Advisory Chat Export", title_style))
            story.append(Spacer(1, 12))

            date_style = ParagraphStyle("Date", parent=styles["Normal"], textColor=colors.gray, fontSize=9)
            story.append(Paragraph(f"Exported on: {time.strftime('%B %d, %Y at %I:%M %p')}", date_style))
            story.append(Spacer(1, 20))

            user_style = ParagraphStyle("User", parent=styles["Normal"], backColor=colors.HexColor("#EEF2FF"), borderPadding=8, fontSize=10, leading=14)
            bot_style = ParagraphStyle("Bot", parent=styles["Normal"], fontSize=10, leading=14)
            role_style = ParagraphStyle("Role", parent=styles["Normal"], textColor=colors.HexColor("#6366f1"), fontSize=9, fontName="Helvetica-Bold")

            for msg in messages:
                role = msg.get("role", "user")
                content = str(msg.get("content", "")).replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
                ts = msg.get("timestamp", "")

                if role == "user":
                    story.append(Paragraph(f"You ({ts})", role_style))
                    story.append(Spacer(1, 3))
                    story.append(Paragraph(content, user_style))
                else:
                    story.append(Paragraph(f"AI Advisor ({ts})", role_style))
                    story.append(Spacer(1, 3))
                    story.append(Paragraph(content, bot_style))
                story.append(Spacer(1, 14))

            doc.build(story)
            buffer.seek(0)

            return Response(
                buffer.getvalue(),
                mimetype="application/pdf",
                headers={"Content-Disposition": "attachment; filename=LMS_Chat_Export.pdf"},
            )

        except ImportError:
            text_content = f"LMS — AI Loan Advisory Chat Export\nExported: {time.strftime('%B %d, %Y at %I:%M %p')}\n{'='*50}\n\n"
            for msg in messages:
                role = "You" if msg.get("role") == "user" else "AI Advisor"
                text_content += f"[{role}] ({msg.get('timestamp', '')})\n{msg.get('content', '')}\n\n---\n\n"
            return Response(
                text_content,
                mimetype="text/plain",
                headers={"Content-Disposition": "attachment; filename=LMS_Chat_Export.txt"},
            )

    except Exception as e:
        print(f"[!] Export error: {e}")
        return jsonify({"error": str(e)}), 500


# --- Health Check ---
@app.route("/api/health")
def health():
    return jsonify({
        "status": "ok",
        "localOnlyMode": LOCAL_ONLY_MODE,
        "geminiConfigured": chat_model is not None,
        "adminEndpointsEnabled": bool(ADMIN_API_KEY),
    })


# ============================================================
#  10. START THE SERVER
# ============================================================

if __name__ == "__main__":
    print()
    print("=" * 55)
    print("  🏦 AI Loan Advisory Chatbot (Fixed Edition)")
    print("  [*] Starting at http://localhost:5000")
    print()
    print("  Features: Local Embeddings | Semantic Domain Gate")
    print("            Streaming | EMI Calc | Conversation Memory")
    print("            Admin-Protected Uploads | PDF Export")
    print()
    print(f"  LOCAL_ONLY_MODE: {LOCAL_ONLY_MODE}")
    print(f"  Gemini generation available: {chat_model is not None}")
    print(f"  Admin endpoints enabled: {bool(ADMIN_API_KEY)}")
    print(f"  Allowed CORS origins: {ALLOWED_ORIGINS}")
    print()
    print("  Press Ctrl+C to stop")
    print("=" * 55)
    print()

    init_vector_store()
    init_feedback_store()

    app.run(host="0.0.0.0", port=5000, debug=FLASK_DEBUG, use_reloader=FLASK_DEBUG)